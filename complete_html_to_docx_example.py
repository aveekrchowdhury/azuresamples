#!/usr/bin/env python3
"""
Complete HTML to DOCX Conversion Example

This script demonstrates how to convert HTML content to Microsoft Word DOCX format.
It includes examples for various HTML elements and formatting options.
"""

import os
import sys
from datetime import datetime

# Import required libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup
    import html
    print("✅ All required packages imported successfully!")
except ImportError as e:
    print(f"❌ Missing package: {e}")
    print("Please install: pip install python-docx beautifulsoup4 lxml")
    sys.exit(1)

def html_to_docx_converter(html_content, output_filename=None):
    """
    Convert HTML content to DOCX format
    
    Args:
        html_content (str): HTML content as string
        output_filename (str): Optional output filename
        
    Returns:
        str: Path to created DOCX file
    """
    
    print("🔄 Converting HTML to DOCX...")
    
    # Create a new document
    doc = Document()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove script and style tags
    for script in soup(["script", "style"]):
        script.decompose()
    
    # Add title if present
    title = soup.find('title')
    if title and title.get_text().strip():
        doc.add_heading(title.get_text().strip(), 0)
    
    # Process body content
    body = soup.find('body') or soup
    
    for element in body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'div']):
        process_element(doc, element)
    
    # Generate filename if not provided
    if not output_filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"converted_html_{timestamp}.docx"
    
    # Save document
    doc.save(output_filename)
    return output_filename

def process_element(doc, element):
    """Process individual HTML elements"""
    
    tag_name = element.name.lower()
    
    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        # Handle headings
        level = int(tag_name[1])
        heading = doc.add_heading(level=level)
        add_formatted_text(heading, element)
        
    elif tag_name == 'p':
        # Handle paragraphs
        if element.get_text().strip():  # Only add non-empty paragraphs
            paragraph = doc.add_paragraph()
            add_formatted_text(paragraph, element)
            
    elif tag_name in ['ul', 'ol']:
        # Handle lists
        process_list(doc, element)
        
    elif tag_name == 'table':
        # Handle tables
        process_table(doc, element)
        
    elif tag_name == 'div':
        # Handle divs by processing their children
        for child in element.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table'], recursive=False):
            process_element(doc, child)

def add_formatted_text(paragraph, element):
    """Add formatted text to a paragraph"""
    
    for content in element.contents:
        if hasattr(content, 'name') and content.name:
            # HTML element
            tag = content.name.lower()
            text = content.get_text()
            
            if tag in ['strong', 'b']:
                run = paragraph.add_run(text)
                run.bold = True
            elif tag in ['em', 'i']:
                run = paragraph.add_run(text)
                run.italic = True
            elif tag == 'u':
                run = paragraph.add_run(text)
                run.underline = True
            elif tag == 'a':
                url = content.get('href', '')
                link_text = f"{text} ({url})" if url else text
                run = paragraph.add_run(link_text)
            elif tag == 'code':
                run = paragraph.add_run(text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                paragraph.add_run(text)
        else:
            # Text content
            if str(content).strip():
                paragraph.add_run(str(content))

def process_list(doc, list_element):
    """Process HTML lists"""
    
    list_items = list_element.find_all('li', recursive=False)
    list_style = 'List Number' if list_element.name == 'ol' else 'List Bullet'
    
    for item in list_items:
        paragraph = doc.add_paragraph(style=list_style)
        add_formatted_text(paragraph, item)

def process_table(doc, table_element):
    """Process HTML tables"""
    
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Count columns
    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            if col_idx < max_cols:
                table_cell = table.cell(row_idx, col_idx)
                # Clear default text
                table_cell.text = ''
                # Add content
                paragraph = table_cell.paragraphs[0]
                add_formatted_text(paragraph, cell)
                
                # Make headers bold
                if cell.name == 'th':
                    for run in paragraph.runs:
                        run.bold = True

def main():
    """Main function with examples"""
    
    print("🌟 HTML to DOCX Converter")
    print("=" * 50)
    
    # Example 1: Simple HTML
    simple_html = """
    <html>
    <head><title>Simple Document</title></head>
    <body>
        <h1>Welcome</h1>
        <p>This is a <strong>simple</strong> HTML document with <em>basic</em> formatting.</p>
    </body>
    </html>
    """
    
    print("📝 Example 1: Converting simple HTML...")
    output1 = html_to_docx_converter(simple_html, "example1_simple.docx")
    print(f"   ✅ Created: {output1}")
    
    # Example 2: Complex HTML with tables and lists
    complex_html = """
    <html>
    <head><title>Complex Document</title></head>
    <body>
        <h1>Project Report</h1>
        <p>This report contains <strong>important</strong> information about our project.</p>
        
        <h2>Team Members</h2>
        <ul>
            <li><strong>John Doe</strong> - Project Manager</li>
            <li><em>Jane Smith</em> - Developer</li>
            <li><u>Bob Johnson</u> - Designer</li>
        </ul>
        
        <h2>Project Timeline</h2>
        <ol>
            <li>Planning Phase</li>
            <li>Development Phase</li>
            <li>Testing Phase</li>
            <li>Deployment Phase</li>
        </ol>
        
        <h2>Budget Summary</h2>
        <table>
            <tr>
                <th>Category</th>
                <th>Amount</th>
                <th>Status</th>
            </tr>
            <tr>
                <td>Development</td>
                <td>$50,000</td>
                <td>Approved</td>
            </tr>
            <tr>
                <td>Testing</td>
                <td>$15,000</td>
                <td>Pending</td>
            </tr>
            <tr>
                <td>Deployment</td>
                <td>$10,000</td>
                <td>Approved</td>
            </tr>
        </table>
        
        <h3>Notes</h3>
        <p>For more information, visit our website at <a href="https://example.com">example.com</a>.</p>
        <p>Code snippet: <code>print("Hello, World!")</code></p>
    </body>
    </html>
    """
    
    print("📝 Example 2: Converting complex HTML...")
    output2 = html_to_docx_converter(complex_html, "example2_complex.docx")
    print(f"   ✅ Created: {output2}")
    
    # Show file information
    print(f"\n📊 File Information:")
    for filename in [output1, output2]:
        if os.path.exists(filename):
            size = os.path.getsize(filename)
            print(f"   📄 {filename}: {size:,} bytes")
    
    print(f"\n🎉 Conversion completed!")
    print(f"   💡 You can now open the generated DOCX files to view the results.")
    
    # Usage instructions
    print(f"\n💡 USAGE INSTRUCTIONS:")
    print(f"   1. Import the converter function")
    print(f"   2. Pass your HTML content as a string")
    print(f"   3. Specify an output filename (optional)")
    print(f"   4. The function returns the path to the created DOCX file")
    
    print(f"\n📋 Example Code:")
    print(f'   output_file = html_to_docx_converter(html_content, "my_document.docx")')

if __name__ == "__main__":
    main()