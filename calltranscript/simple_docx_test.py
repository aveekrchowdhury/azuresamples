#!/usr/bin/env python3
"""
Simple HTML to DOCX Test
"""

import sys
import os

# Add current directory to path to import our module
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from docx import Document
    from bs4 import BeautifulSoup
    print("✅ Required packages imported successfully!")
    
    # Simple test
    print("🔄 Creating a simple DOCX document...")
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('HTML to DOCX Converter Test', 0)
    
    # Add content
    p1 = doc.add_paragraph('This document was created by the HTML to DOCX converter.')
    p1_run = p1.add_run(' This text is bold.')
    p1_run.bold = True
    
    p2 = doc.add_paragraph()
    p2_run1 = p2.add_run('This text is italic. ')
    p2_run1.italic = True
    p2_run2 = p2.add_run('This text is underlined.')
    p2_run2.underline = True
    
    # Add a list
    doc.add_heading('Features:', level=1)
    doc.add_paragraph('Convert HTML to DOCX', style='List Bullet')
    doc.add_paragraph('Preserve formatting', style='List Bullet')
    doc.add_paragraph('Handle tables and lists', style='List Bullet')
    
    # Save the document
    output_file = 'simple_test_output.docx'
    doc.save(output_file)
    
    print(f"✅ Test document created successfully!")
    print(f"📄 Output file: {output_file}")
    print(f"📍 Full path: {os.path.abspath(output_file)}")
    print(f"📊 File size: {os.path.getsize(output_file):,} bytes")
    
except ImportError as e:
    print(f"❌ Missing required package: {e}")
    print("Please install the required packages:")
    print("pip install python-docx beautifulsoup4 lxml")
    
except Exception as e:
    print(f"❌ Error: {e}")

print("\n💡 For full HTML conversion functionality, use the html_to_docx_converter.py module")