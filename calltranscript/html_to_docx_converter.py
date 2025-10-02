# HTML to DOCX Converter
# This script converts HTML content to a Microsoft Word DOCX file

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
import html
import os
from datetime import datetime

class HTMLToDOCXConverter:
    """
    A class to convert HTML content to DOCX format
    """
    
    def __init__(self):
        self.document = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.document.styles
        
        # Create heading styles if they don't exist
        for level in range(1, 7):
            style_name = f'Custom Heading {level}'
            if style_name not in styles:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                heading_style.base_style = styles['Normal']
                font = heading_style.font
                font.name = 'Calibri'
                font.size = Pt(18 - level)
                font.bold = True
                
        # Create emphasis styles
        if 'Custom Bold' not in styles:
            bold_style = styles.add_style('Custom Bold', WD_STYLE_TYPE.CHARACTER)
            bold_style.font.bold = True
            
        if 'Custom Italic' not in styles:
            italic_style = styles.add_style('Custom Italic', WD_STYLE_TYPE.CHARACTER)
            italic_style.font.italic = True
    
    def clean_html(self, html_content):
        """Clean and normalize HTML content"""
        # Decode HTML entities
        html_content = html.unescape(html_content)
        
        # Parse with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        return soup
    
    def add_paragraph_with_formatting(self, element, parent_paragraph=None):
        """Add a paragraph with proper formatting based on HTML element"""
        
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Handle headings
            level = int(element.name[1])
            heading = self.document.add_heading(level=level)
            heading.text = element.get_text().strip()
            
        elif element.name == 'p':
            # Handle paragraphs
            paragraph = self.document.add_paragraph()
            self.process_inline_elements(element, paragraph)
            
        elif element.name in ['ul', 'ol']:
            # Handle lists
            self.process_list(element)
            
        elif element.name == 'table':
            # Handle tables
            self.process_table(element)
            
        elif element.name == 'br':
            # Handle line breaks
            self.document.add_paragraph()
            
        elif element.name in ['div', 'section', 'article']:
            # Handle block elements by processing their children
            for child in element.children:
                if hasattr(child, 'name') and child.name:
                    self.add_paragraph_with_formatting(child)
                elif child.string and child.string.strip():
                    # Handle text nodes
                    paragraph = self.document.add_paragraph()
                    paragraph.add_run(child.string.strip())
    
    def process_inline_elements(self, element, paragraph):
        """Process inline elements within a paragraph"""
        for child in element.children:
            if hasattr(child, 'name') and child.name:
                if child.name == 'strong' or child.name == 'b':
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'u':
                    run = paragraph.add_run(child.get_text())
                    run.underline = True
                elif child.name == 'a':
                    # Handle links
                    link_text = child.get_text()
                    url = child.get('href', '')
                    run = paragraph.add_run(f"{link_text} ({url})")
                    run.font.color.rgb = None  # Blue color for links
                elif child.name == 'code':
                    run = paragraph.add_run(child.get_text())
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    # Recursively process other inline elements
                    self.process_inline_elements(child, paragraph)
            elif child.string:
                # Handle text nodes
                paragraph.add_run(child.string)
    
    def process_list(self, list_element):
        """Process HTML lists (ul/ol)"""
        is_ordered = list_element.name == 'ol'
        
        for li in list_element.find_all('li', recursive=False):
            paragraph = self.document.add_paragraph()
            
            if is_ordered:
                paragraph.style = 'List Number'
            else:
                paragraph.style = 'List Bullet'
            
            # Process the content of the list item
            self.process_inline_elements(li, paragraph)
    
    def process_table(self, table_element):
        """Process HTML tables"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Count maximum columns
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        
        # Create table
        table = self.document.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    table_cell = table.cell(i, j)
                    # Clear the cell first
                    table_cell.text = ''
                    # Add content
                    paragraph = table_cell.paragraphs[0]
                    self.process_inline_elements(cell, paragraph)
                    
                    # Make header cells bold
                    if cell.name == 'th':
                        for run in paragraph.runs:
                            run.bold = True
    
    def convert_html_to_docx(self, html_content, output_path=None):
        """
        Convert HTML content to DOCX file
        
        Args:
            html_content (str): HTML content to convert
            output_path (str): Path for output DOCX file
            
        Returns:
            str: Path to the created DOCX file
        """
        
        # Clean the HTML
        soup = self.clean_html(html_content)
        
        # Add document title if HTML has a title
        title_element = soup.find('title')
        if title_element:
            title = title_element.get_text().strip()
            if title:
                heading = self.document.add_heading(level=0)
                heading.text = title
        
        # Process the body content
        body = soup.find('body')
        if body:
            # Process each top-level element in body
            for element in body.children:
                if hasattr(element, 'name') and element.name:
                    self.add_paragraph_with_formatting(element)
                elif element.string and element.string.strip():
                    # Handle text nodes at body level
                    paragraph = self.document.add_paragraph()
                    paragraph.add_run(element.string.strip())
        else:
            # If no body tag, process the entire soup
            for element in soup.children:
                if hasattr(element, 'name') and element.name:
                    self.add_paragraph_with_formatting(element)
        
        # Generate output path if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"converted_document_{timestamp}.docx"
        
        # Save the document
        self.document.save(output_path)
        
        return output_path

def convert_html_file_to_docx(html_file_path, output_path=None):
    """
    Convert an HTML file to DOCX
    
    Args:
        html_file_path (str): Path to the HTML file
        output_path (str): Path for output DOCX file
        
    Returns:
        str: Path to the created DOCX file
    """
    
    if not os.path.exists(html_file_path):
        raise FileNotFoundError(f"HTML file not found: {html_file_path}")
    
    with open(html_file_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    
    converter = HTMLToDOCXConverter()
    return converter.convert_html_to_docx(html_content, output_path)

def convert_html_string_to_docx(html_string, output_path=None):
    """
    Convert an HTML string to DOCX
    
    Args:
        html_string (str): HTML content as string
        output_path (str): Path for output DOCX file
        
    Returns:
        str: Path to the created DOCX file
    """
    
    converter = HTMLToDOCXConverter()
    return converter.convert_html_to_docx(html_string, output_path)

# Example usage and testing
if __name__ == "__main__":
    # Sample HTML content for testing
    sample_html = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Sample Document</title>
    </head>
    <body>
        <h1>Main Heading</h1>
        <p>This is a <strong>bold</strong> paragraph with <em>italic</em> text and a <u>underlined</u> word.</p>
        
        <h2>Subheading</h2>
        <p>This paragraph contains a <a href="https://example.com">link to example.com</a>.</p>
        
        <h3>Lists Example</h3>
        <ul>
            <li>First bullet point</li>
            <li>Second bullet point with <strong>bold text</strong></li>
            <li>Third bullet point</li>
        </ul>
        
        <ol>
            <li>First numbered item</li>
            <li>Second numbered item</li>
            <li>Third numbered item</li>
        </ol>
        
        <h3>Table Example</h3>
        <table>
            <tr>
                <th>Header 1</th>
                <th>Header 2</th>
                <th>Header 3</th>
            </tr>
            <tr>
                <td>Cell 1,1</td>
                <td>Cell 1,2</td>
                <td>Cell 1,3</td>
            </tr>
            <tr>
                <td>Cell 2,1</td>
                <td>Cell 2,2</td>
                <td>Cell 2,3</td>
            </tr>
        </table>
        
        <p>This is a paragraph with <code>inline code</code> formatting.</p>
        
        <div>
            <p>This is a paragraph inside a div.</p>
            <p>Another paragraph with mixed formatting: <strong><em>bold and italic</em></strong>.</p>
        </div>
    </body>
    </html>
    """
    
    print("🔄 Converting sample HTML to DOCX...")
    
    try:
        # Convert the sample HTML
        output_file = convert_html_string_to_docx(sample_html, "sample_converted_document.docx")
        print(f"✅ Conversion successful!")
        print(f"📄 Output file: {output_file}")
        print(f"📍 Full path: {os.path.abspath(output_file)}")
        
    except Exception as e:
        print(f"❌ Conversion failed: {str(e)}")