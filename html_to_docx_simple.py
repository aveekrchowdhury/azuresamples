#!/usr/bin/env python3
"""
HTML to DOCX Converter - Simple and Reliable Version

This script provides a straightforward way to convert HTML content to DOCX format.
"""

import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from bs4 import BeautifulSoup
    print("✅ Required packages loaded successfully!")
except ImportError as e:
    print(f"❌ Missing package: {e}")
    print("Install with: pip install python-docx beautifulsoup4 lxml")
    sys.exit(1)

def convert_html_to_docx(html_content, output_filename=None):
    """
    Convert HTML content to DOCX format
    
    Args:
        html_content (str): HTML content as string
        output_filename (str): Optional output filename
        
    Returns:
        str: Path to created DOCX file
    """
    
    # Create new document
    doc = Document()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Add title if present
    title = soup.find('title')
    if title and title.string:
        doc.add_heading(title.string.strip(), 0)
    
    # Process content
    body = soup.find('body') if soup.find('body') else soup
    
    # Handle different elements
    for element in body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
        level = int(element.name[1])
        doc.add_heading(element.get_text().strip(), level)
    
    for element in body.find_all('p'):
        text = element.get_text().strip()
        if text:
            paragraph = doc.add_paragraph()
            
            # Handle basic formatting
            for child in element.children:
                if hasattr(child, 'name') and child.name:
                    if child.name in ['strong', 'b']:
                        run = paragraph.add_run(child.get_text())
                        run.bold = True
                    elif child.name in ['em', 'i']:
                        run = paragraph.add_run(child.get_text())
                        run.italic = True
                    elif child.name == 'u':
                        run = paragraph.add_run(child.get_text())
                        run.underline = True
                    else:
                        paragraph.add_run(child.get_text())
                else:
                    if str(child).strip():
                        paragraph.add_run(str(child))
    
    # Handle lists
    for ul in body.find_all('ul'):
        for li in ul.find_all('li'):
            doc.add_paragraph(li.get_text().strip(), style='List Bullet')
    
    for ol in body.find_all('ol'):
        for li in ol.find_all('li'):
            doc.add_paragraph(li.get_text().strip(), style='List Number')
    
    # Handle tables
    for table in body.find_all('table'):
        rows = table.find_all('tr')
        if rows:
            # Count columns
            max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            
            # Create table
            doc_table = doc.add_table(rows=len(rows), cols=max_cols)
            doc_table.style = 'Table Grid'
            
            for row_idx, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for col_idx, cell in enumerate(cells):
                    if col_idx < max_cols:
                        doc_table.cell(row_idx, col_idx).text = cell.get_text().strip()
    
    # Generate filename if not provided
    if not output_filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"html_to_docx_{timestamp}.docx"
    
    # Save document
    doc.save(output_filename)
    return output_filename

def main():
    """Demonstration of HTML to DOCX conversion"""
    
    print("🌟 HTML to DOCX Converter Demo")
    print("=" * 40)
    
    # Sample HTML content
    sample_html = """
    <html>
    <head>
        <title>Sample Report</title>
    </head>
    <body>
        <h1>Project Status Report</h1>
        <p>This is the <strong>quarterly report</strong> for our project.</p>
        
        <h2>Key Achievements</h2>
        <ul>
            <li>Completed phase 1 development</li>
            <li>Successfully tested core features</li>
            <li>Deployed to staging environment</li>
        </ul>
        
        <h2>Next Steps</h2>
        <ol>
            <li>User acceptance testing</li>
            <li>Performance optimization</li>
            <li>Production deployment</li>
        </ol>
        
        <h2>Team Performance</h2>
        <table>
            <tr>
                <th>Team Member</th>
                <th>Role</th>
                <th>Status</th>
            </tr>
            <tr>
                <td>Alice Johnson</td>
                <td>Developer</td>
                <td>Excellent</td>
            </tr>
            <tr>
                <td>Bob Smith</td>
                <td>Designer</td>
                <td>Good</td>
            </tr>
        </table>
        
        <p>For questions, contact the <em>project manager</em>.</p>
    </body>
    </html>
    """
    
    print("🔄 Converting HTML to DOCX...")
    
    try:
        output_file = convert_html_to_docx(sample_html, "sample_report.docx")
        
        print(f"✅ Conversion successful!")
        print(f"📄 Output file: {output_file}")
        print(f"📍 Full path: {os.path.abspath(output_file)}")
        
        if os.path.exists(output_file):
            file_size = os.path.getsize(output_file)
            print(f"📊 File size: {file_size:,} bytes")
        
        print(f"\n💡 Usage Examples:")
        print(f"   # Convert HTML string")
        print(f"   output = convert_html_to_docx(html_content)")
        print(f"   ")
        print(f"   # Convert with custom filename")
        print(f"   output = convert_html_to_docx(html_content, 'my_document.docx')")
        
    except Exception as e:
        print(f"❌ Error during conversion: {e}")

if __name__ == "__main__":
    main()